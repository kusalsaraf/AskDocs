from __future__ import annotations

import io
import time

from apps.core.logging import get_logger
from apps.documents.exceptions import ParserProviderError, UnsupportedFileType
from apps.documents.parsing.base import ParsedDocument, ParsedElement, ParserProvider

logger = get_logger(__name__)


class SimpleParserProvider(ParserProvider):
    """Lightweight DOCX + TXT parser — used as fallback when unstructured isn't active."""

    provider_name = "simple"
    supported_file_types = {"docx", "txt"}
    supported_strategies = {"fast"}

    def parse(
        self,
        file_bytes: bytes,
        file_type: str,
        strategy: str = "fast",
    ) -> ParsedDocument:
        if file_type == "txt":
            return self._parse_txt(file_bytes)
        if file_type == "docx":
            return self._parse_docx(file_bytes)
        raise UnsupportedFileType(file_type)

    def _parse_txt(self, file_bytes: bytes) -> ParsedDocument:
        start = time.monotonic()
        try:
            text = file_bytes.decode("utf-8", errors="replace").strip()
        except Exception as exc:
            raise ParserProviderError(f"Failed to decode TXT: {exc}") from exc

        duration_ms = int((time.monotonic() - start) * 1000)
        elements = (
            [ParsedElement(content=text, element_type="NarrativeText", page_number=1)]
            if text
            else []
        )
        logger.info(
            "simple txt parse complete",
            extra={"num_elements": len(elements), "duration_ms": duration_ms},
        )
        return ParsedDocument(elements=elements, metadata={"page_count": 1})

    def _parse_docx(self, file_bytes: bytes) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ParserProviderError("python-docx is not installed.") from exc

        start = time.monotonic()
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            elements: list[ParsedElement] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name if para.style else ""
                if "Heading" in style_name or "Title" in style_name:
                    element_type = "Title"
                elif "List" in style_name:
                    element_type = "ListItem"
                else:
                    element_type = "NarrativeText"
                elements.append(
                    ParsedElement(
                        content=text,
                        element_type=element_type,
                        page_number=None,
                    )
                )
        except Exception as exc:
            raise ParserProviderError(f"python-docx failed: {exc}") from exc

        duration_ms = int((time.monotonic() - start) * 1000)
        logger.info(
            "simple docx parse complete",
            extra={"num_elements": len(elements), "duration_ms": duration_ms},
        )
        return ParsedDocument(elements=elements, metadata={"page_count": None})
